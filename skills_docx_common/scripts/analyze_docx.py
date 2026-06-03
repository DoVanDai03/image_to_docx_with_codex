import argparse
import sys
import zipfile
from collections import Counter
from pathlib import Path

from lxml import etree

from common import DEFAULT_SRC_COMMON, get_common_paths, require_file, resolve_path, write_json


def import_document():
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "Missing dependency python-docx. Run: python -m pip install -r "
            "skills_docx_common/requirements.txt"
        ) from exc
    return Document


def local_name(tag: str) -> str:
    return etree.QName(tag).localname if tag else ""


def read_package_analysis(docx_path: Path) -> dict:
    package = {
        "has_document_xml": False,
        "media": [],
        "relationships": [],
        "body_children": [],
        "element_counts": {},
    }
    with zipfile.ZipFile(docx_path, "r") as zf:
        names = zf.namelist()
        package["media"] = [name for name in names if name.startswith("word/media/")]
        package["relationships"] = [name for name in names if name.endswith(".rels")]
        if "word/document.xml" not in names:
            return package

        package["has_document_xml"] = True
        root = etree.fromstring(zf.read("word/document.xml"))
        body = root.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body")
        if body is not None:
            package["body_children"] = [local_name(child.tag) for child in body]
        counts = Counter(local_name(el.tag) for el in root.iter())
        package["element_counts"] = dict(sorted(counts.items()))
    return package


def analyze_docx(docx_path: Path) -> dict:
    Document = import_document()
    doc = Document(str(docx_path))

    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "text": text,
                "style": paragraph.style.name if paragraph.style else "",
                "runs": len(paragraph.runs),
            }
        )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "header": rows[0] if rows else [],
                "sample_rows": rows[:5],
            }
        )

    sections = []
    for section_index, section in enumerate(doc.sections):
        sections.append(
            {
                "index": section_index,
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )

    return {
        "source": str(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "non_empty_paragraphs": paragraphs,
        "table_count": len(doc.tables),
        "tables": tables,
        "sections": sections,
        "package": read_package_analysis(docx_path),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Analyze a common DOCX template.")
    parser.add_argument("--docx", default="", help="DOCX path.")
    parser.add_argument("--out", default="", help="Output JSON path.")
    args = parser.parse_args(argv)

    default_docx = get_common_paths(DEFAULT_SRC_COMMON)["docx"]
    docx_path = require_file(resolve_path(args.docx, default_docx), "DOCX common")
    out_path = resolve_path(args.out, docx_path.with_suffix(".analysis.json"))
    write_json(out_path, analyze_docx(docx_path))
    print(f"docx_analysis: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
