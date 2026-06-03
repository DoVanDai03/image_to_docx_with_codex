import argparse
import re
import sys
from pathlib import Path

from common import (
    DEFAULT_SRC_COMMON,
    get_common_paths,
    load_json,
    normalize_workflow,
    require_file,
    resolve_path,
    slugify,
)


def import_document():
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Emu, Pt, Twips
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "Missing dependency python-docx. Run: python -m pip install -r "
            "docx_common_workflow/skills_docx_common/requirements.txt"
        ) from exc
    return {
        "Document": Document,
        "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_LINE_SPACING": WD_LINE_SPACING,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Emu": Emu,
        "Pt": Pt,
        "Twips": Twips,
    }


DOCX = None
FONT_BODY = "Times New Roman"
FONT_HEADING = "Times New Roman"
SIZE_PT = 12
NUMBER_PREFIX_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)?|A\d+)\s*[\.)]\s+")


def clear_document_body(doc) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def configure_page(doc) -> None:
    twips = DOCX["Twips"]
    for section in doc.sections:
        section.page_width = twips(12240)
        section.page_height = twips(15840)
        section.top_margin = twips(1440)
        section.right_margin = twips(1440)
        section.bottom_margin = twips(1440)
        section.left_margin = twips(1440)
        section.header_distance = twips(720)
        section.footer_distance = twips(720)


def style_font_name(style, fallback: str = "") -> str:
    if style is None:
        return fallback
    if style.font.name:
        return style.font.name
    r_pr = style._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return fallback
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        value = r_pr.rFonts.get(DOCX["qn"](attr))
        if value:
            return value
    return fallback


def is_legacy_vn_font(font_name: str) -> bool:
    return str(font_name or "").startswith(".Vn")


def configure_defaults(doc) -> None:
    """Keep document defaults from docx_common.docx as the source of truth."""
    global FONT_BODY, FONT_HEADING
    try:
        normal_style = doc.styles["Normal"]
    except KeyError:
        normal_style = None

    FONT_BODY = style_font_name(normal_style, FONT_BODY)

    heading_style = None
    for style_name in ("Heading 7", "Heading 6", "Heading 5"):
        try:
            heading_style = doc.styles[style_name]
            break
        except KeyError:
            continue
    heading_font = style_font_name(heading_style, FONT_BODY)
    FONT_HEADING = FONT_BODY if is_legacy_vn_font(heading_font) else heading_font


def disable_heading_auto_numbering(doc) -> None:
    """Generated headings carry manual section numbers, so suppress style numbering."""
    for style_name in ("Heading 6", "Heading 7"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        p_pr = style._element.pPr
        if p_pr is None:
            continue
        num_pr = p_pr.find(DOCX["qn"]("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)


def ensure_child(parent, tag):
    child = parent.find(DOCX["qn"](tag))
    if child is None:
        child = DOCX["OxmlElement"](tag)
        parent.append(child)
    return child


def set_paragraph_lang(paragraph, val="vi-VN", east_asia="vi-VN") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        lang = ensure_child(r_pr, "w:lang")
        lang.set(DOCX["qn"]("w:val"), val)
        lang.set(DOCX["qn"]("w:eastAsia"), east_asia)


def set_run_font(run, font_name=FONT_BODY, size_pt=SIZE_PT, bold=False, italic=False, underline=False) -> None:
    run.font.name = font_name
    run.font.size = DOCX["Pt"](size_pt)
    run.bold = True if bold else None
    run.italic = True if italic else None
    run.underline = True if underline else None
    run._element.rPr.rFonts.set(DOCX["qn"]("w:eastAsia"), font_name)


def set_paragraph_format(paragraph, alignment=None, before=0, after=160, line=259, keep=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = DOCX["Twips"](before)
    pf.space_after = DOCX["Twips"](after)
    pf.line_spacing_rule = DOCX["WD_LINE_SPACING"].MULTIPLE
    pf.line_spacing = 1.15 if line else None
    if alignment is not None:
        paragraph.alignment = alignment
    if keep:
        p_pr = paragraph._p.get_or_add_pPr()
        p_pr.append(DOCX["OxmlElement"]("w:keepNext"))
        p_pr.append(DOCX["OxmlElement"]("w:keepLines"))


def set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(DOCX["qn"]("w:outlineLvl"))
    if outline is None:
        outline = DOCX["OxmlElement"]("w:outlineLvl")
        p_pr.append(outline)
    outline.set(DOCX["qn"]("w:val"), str(level))


def make_heading(doc, number_text: str, title_text: str, outline_level: int = 5):
    paragraph = doc.add_paragraph()
    style_name = "Heading 6" if outline_level <= 4 else "Heading 7"
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        pass
    set_paragraph_format(paragraph, before=40, after=160, keep=True)
    set_outline_level(paragraph, outline_level)
    number = paragraph.add_run(number_text)
    set_run_font(number, FONT_HEADING, SIZE_PT, italic=True)
    title = paragraph.add_run(title_text)
    set_run_font(title, FONT_HEADING, SIZE_PT, italic=True, underline=True)
    set_paragraph_lang(paragraph, val="pt-BR", east_asia="ja-JP")
    return paragraph


def set_cell_width(cell, width: int, width_type: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(DOCX["qn"]("w:tcW"))
    if tc_w is None:
        tc_w = DOCX["OxmlElement"]("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(DOCX["qn"]("w:w"), str(width))
    tc_w.set(DOCX["qn"]("w:type"), width_type)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(DOCX["qn"]("w:shd"))
    if shd is None:
        shd = DOCX["OxmlElement"]("w:shd")
        tc_pr.append(shd)
    shd.set(DOCX["qn"]("w:fill"), fill)
    shd.set(DOCX["qn"]("w:val"), "clear")


def set_table_width(table, width: int, width_type: str) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(DOCX["qn"]("w:tblW"))
    if tbl_w is None:
        tbl_w = DOCX["OxmlElement"]("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(DOCX["qn"]("w:w"), str(width))
    tbl_w.set(DOCX["qn"]("w:type"), width_type)


def set_table_indent(table, width: int) -> None:
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.find(DOCX["qn"]("w:tblInd"))
    if indent is None:
        indent = DOCX["OxmlElement"]("w:tblInd")
        tbl_pr.append(indent)
    indent.set(DOCX["qn"]("w:w"), str(width))
    indent.set(DOCX["qn"]("w:type"), "dxa")


def set_table_grid(table, widths: list[int]) -> None:
    tbl = table._tbl
    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = DOCX["OxmlElement"]("w:tblGrid")
    for width in widths:
        col = DOCX["OxmlElement"]("w:gridCol")
        col.set(DOCX["qn"]("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)


def set_table_look(table, value: str, first_row: str = "0", last_row: str = "0", first_column: str = "0", last_column: str = "0") -> None:
    tbl_pr = table._tbl.tblPr
    look = tbl_pr.find(DOCX["qn"]("w:tblLook"))
    if look is None:
        look = DOCX["OxmlElement"]("w:tblLook")
        tbl_pr.append(look)
    look.set(DOCX["qn"]("w:val"), value)
    look.set(DOCX["qn"]("w:firstRow"), first_row)
    look.set(DOCX["qn"]("w:lastRow"), last_row)
    look.set(DOCX["qn"]("w:firstColumn"), first_column)
    look.set(DOCX["qn"]("w:lastColumn"), last_column)
    look.set(DOCX["qn"]("w:noHBand"), "0")
    look.set(DOCX["qn"]("w:noVBand"), "0")


def set_table_borders(table, style: str, size: int, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(DOCX["qn"]("w:tblBorders"))
    if borders is None:
        borders = DOCX["OxmlElement"]("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(DOCX["qn"](tag))
        if element is None:
            element = DOCX["OxmlElement"](tag)
            borders.append(element)
        element.set(DOCX["qn"]("w:val"), style)
        element.set(DOCX["qn"]("w:sz"), str(size))
        element.set(DOCX["qn"]("w:space"), "0")
        element.set(DOCX["qn"]("w:color"), color)


def set_row_height(row, height: int, exact: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = tr_pr.find(DOCX["qn"]("w:trHeight"))
    if tr_h is None:
        tr_h = DOCX["OxmlElement"]("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(DOCX["qn"]("w:val"), str(height))
    if exact:
        tr_h.set(DOCX["qn"]("w:hRule"), "exact")
    elif tr_h.get(DOCX["qn"]("w:hRule")) is not None:
        del tr_h.attrib[DOCX["qn"]("w:hRule")]


def set_cell_text(cell, text: str, *, bold=False, alignment=None, indent_left=0, fill="", width=0, width_type="dxa", after=160):
    cell.text = ""
    if width:
        set_cell_width(cell, width, width_type)
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = DOCX["WD_CELL_VERTICAL_ALIGNMENT"].CENTER
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, alignment=alignment, before=0, after=after)
    paragraph.paragraph_format.left_indent = DOCX["Twips"](indent_left)
    value = "UNKNOWN" if text is None else str(text)
    run = paragraph.add_run(value)
    set_run_font(run, FONT_BODY, SIZE_PT, bold=bold)
    set_paragraph_lang(paragraph)


def make_info_table(doc, workflow: dict):
    alignment = DOCX["WD_ALIGN_PARAGRAPH"].JUSTIFY
    table = doc.add_table(rows=0, cols=2)
    table.alignment = DOCX["WD_TABLE_ALIGNMENT"].CENTER
    set_table_width(table, 4855, "pct")
    set_table_grid(table, [2876, 6168])
    set_table_borders(table, "single", 4, "808080")
    set_table_look(table, "0000")

    rows = [
        ("Tên chức năng", workflow["name"]),
        ("Mô tả", workflow.get("purpose", "UNKNOWN")),
        ("Tác nhân", workflow.get("actor", "Người dùng hệ thống")),
        ("Điều kiện trước", workflow.get("preconditions", "Người dùng có quyền thực hiện chức năng.")),
        (
            "Điều kiện sau",
            workflow.get(
                "postconditions",
                "Trường hợp thành công: người dùng thực hiện được chức năng. "
                "Trường hợp không thành công: hệ thống hiển thị thông báo tương ứng.",
            ),
        ),
        ("Ngoại lệ", "; ".join(workflow.get("error_scenarios", [])) or "UNKNOWN"),
        ("Các yêu cầu đặc biệt", "; ".join(workflow.get("business_rules", [])) or "N/A"),
    ]

    row_heights = [284, 284, 395, 378, 530, 284, 284]
    for index, (label, value) in enumerate(rows):
        row = table.add_row()
        set_row_height(row, row_heights[index])
        left, right = row.cells
        set_cell_text(left, label, bold=True, alignment=alignment, indent_left=142, fill="F3F3F3", width=1590, width_type="pct")
        set_cell_text(right, value, alignment=alignment, width=3410, width_type="pct")
    return table


def make_flow_table(doc, rows: list[dict], col_widths: list[int] | None = None, alternative: bool = False):
    col_widths = col_widths or [2572, 4950, 1757]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, 9279, "dxa")
    set_table_indent(table, -34)
    set_table_grid(table, col_widths)
    set_table_borders(table, "dotted", 4, "auto")
    set_table_look(table, "01E0", first_row="1", last_row="1", first_column="1", last_column="1")
    header = table.rows[0]
    set_row_height(header, 530)
    headers = [
        ("Hành động của tác nhân", col_widths[0]),
        ("Phản ứng của hệ thống", col_widths[1]),
        ("Dữ liệu liên quan (C/R/U/D)", col_widths[2]),
    ]
    for cell, (text, width) in zip(header.cells, headers):
        set_cell_text(cell, text, bold=True, alignment=DOCX["WD_ALIGN_PARAGRAPH"].CENTER, fill="F5F5F5", width=width, width_type="dxa", after=0)

    for index, step in enumerate(rows, start=1):
        row = table.add_row()
        actor_text = numbered_flow_text(step.get("actor", "UNKNOWN"), index * 2 - 1)
        system_text = numbered_flow_text(step.get("system", "UNKNOWN"), index * 2)
        set_cell_text(row.cells[0], actor_text, alignment=DOCX["WD_ALIGN_PARAGRAPH"].JUSTIFY, width=col_widths[0], width_type="dxa", after=0)
        set_cell_text(row.cells[1], system_text, alignment=DOCX["WD_ALIGN_PARAGRAPH"].JUSTIFY, width=col_widths[1], width_type="dxa", after=0)
        set_cell_text(row.cells[2], str(step.get("crud", "R")), alignment=DOCX["WD_ALIGN_PARAGRAPH"].CENTER, width=col_widths[2], width_type="dxa", after=0)

    footer = table.add_row()
    for cell, width in zip(footer.cells, col_widths):
        set_cell_text(cell, "", fill="EBEBEB", width=width, width_type="dxa", after=0)
    return table


def numbered_flow_text(value, number: int) -> str:
    text = str(value or "UNKNOWN").strip()
    if NUMBER_PREFIX_RE.match(text):
        return text
    return f"{number}. {text}"


def add_caption(doc, text: str):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, alignment=DOCX["WD_ALIGN_PARAGRAPH"].CENTER, before=0, after=160)
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, SIZE_PT, italic=True)
    return paragraph


def add_flow_image(doc, source_image: str) -> None:
    image_path = Path(source_image) if source_image else None
    if image_path and not image_path.is_absolute():
        image_path = (Path.cwd() / image_path).resolve()

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, alignment=DOCX["WD_ALIGN_PARAGRAPH"].CENTER, before=0, after=160)

    if image_path and image_path.is_file():
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=DOCX["Emu"](5943600))
        return

    run = paragraph.add_run("[Hình biểu đồ luồng]")
    set_run_font(run, FONT_BODY, SIZE_PT)


def add_list_section(doc, title: str, items: list[str], section_number: str):
    make_heading(doc, f"{section_number}. ", title, 5)
    if not items:
        items = ["N/A"]
    for item in items:
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, alignment=DOCX["WD_ALIGN_PARAGRAPH"].JUSTIFY, before=0, after=160)
        run = paragraph.add_run("- " + str(item))
        set_run_font(run, FONT_BODY, SIZE_PT)


def get_alternative_flow_tables(workflow: dict) -> list[list[dict]]:
    tables: list[list[dict]] = []
    for item in workflow.get("alternative_flows", []):
        if isinstance(item, dict):
            steps = list(item.get("steps") or [])
            if steps:
                tables.append(steps)
        elif isinstance(item, list) and item:
            tables.append(item)
    if tables:
        return tables
    fallback = list(workflow.get("alternative_flow", []))
    return [fallback] if fallback else []


def render_docx(template_docx: Path, workflow: dict, output_docx: Path) -> None:
    global DOCX
    DOCX = import_document()
    doc = DOCX["Document"](str(template_docx))
    clear_document_body(doc)
    configure_page(doc)
    configure_defaults(doc)
    disable_heading_auto_numbering(doc)

    section_base = workflow.get("section_number") or workflow.get("sectionNumber") or "4.1.2.3.2"
    make_heading(doc, f"{section_base}. ", workflow["name"], 4)
    make_heading(doc, f"{section_base}.1. ", "Thông tin chung chức năng", 5)
    make_info_table(doc, workflow)

    make_heading(doc, f"{section_base}.2. ", "Biểu đồ luồng chức năng", 5)
    add_flow_image(doc, workflow.get("source_image", ""))
    add_caption(doc, f"Hình 3: Luồng chức năng {workflow['name']}")

    make_heading(doc, f"{section_base}.3. ", "Mô tả dòng sự kiện chính (Basic Flow)", 5)
    make_flow_table(doc, workflow.get("basic_flow", []))

    make_heading(doc, f"{section_base}.4. ", "Mô tả dòng sự kiện phụ (Alternative Flow)", 5)
    alternative_tables = get_alternative_flow_tables(workflow)
    if alternative_tables:
        for rows in alternative_tables:
            make_flow_table(doc, rows, alternative=True)
    else:
        make_flow_table(doc, [], alternative=True)

    notes = list(workflow.get("validation_rules", [])) + list(workflow.get("notes", []))
    add_list_section(doc, "Ghi chú", notes, f"{section_base}.5")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Render DOCX from common DOCX and workflow JSON.")
    parser.add_argument("--workflow-json", required=True)
    parser.add_argument("--docx-common", default="")
    parser.add_argument("--out", default="")
    args = parser.parse_args(argv)

    workflow_path = require_file(Path(args.workflow_json).resolve(), "Workflow JSON")
    workflow = normalize_workflow(load_json(workflow_path))
    default_docx = get_common_paths(DEFAULT_SRC_COMMON)["docx"]
    template_docx = require_file(resolve_path(args.docx_common, default_docx), "DOCX common")
    output_docx = resolve_path(args.out, Path.cwd() / workflow["id"] / f"{slugify(workflow['id'])}.docx")
    render_docx(template_docx, workflow, output_docx)
    print(f"docx: {output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
